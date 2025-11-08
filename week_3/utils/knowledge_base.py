import os
from typing import List, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma


class KnowledgeBase:
    def __init__(
        self,
        embedding_type: str = "openai",
        api_key: Optional[str] = None,
        persist_directory: str = "data/voice_kb/vectorstore"
    ):
        """Initialize knowledge base with embeddings"""
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)
        
        # Initialize embeddings
        if embedding_type == "openai":
            if not api_key:
                raise ValueError("OpenAI API key required")
            self.embeddings = OpenAIEmbeddings(
                model="text-embedding-3-small",
                openai_api_key=api_key
            )
        else:
            self.embeddings = HuggingFaceEmbeddings()
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Load or create vector store
        self.vector_store = None
        self._load_vector_store()
    
    def _load_vector_store(self):
        """Load existing vector store or create new one"""
        try:
            if os.path.exists(self.persist_directory) and os.listdir(self.persist_directory):
                # Check if chroma.sqlite3 exists and fix permissions
                db_path = os.path.join(self.persist_directory, "chroma.sqlite3")
                if os.path.exists(db_path):
                    os.chmod(db_path, 0o666)
                
                self.vector_store = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
        except Exception as e:
            print(f"Could not load existing vector store: {e}. Creating new one...")
            # Delete corrupted database and start fresh
            import shutil
            if os.path.exists(self.persist_directory):
                shutil.rmtree(self.persist_directory)
            os.makedirs(self.persist_directory, exist_ok=True)
            self.vector_store = None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT/MD file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def process_file(self, file_path: str, file_type: str) -> List[Document]:
        """Process file and return documents"""
        # Extract text based on file type
        if file_type == "pdf":
            text = self.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            text = self.extract_text_from_docx(file_path)
        elif file_type in ["txt", "md"]:
            text = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create documents
        file_name = Path(file_path).name
        documents = [
            Document(
                page_content=chunk,
                metadata={"source": file_name, "file_path": file_path}
            )
            for chunk in chunks
        ]
        
        return documents
    
    def add_documents(self, documents: List[Document]):
        """Add documents to vector store"""
        try:
            if self.vector_store is None:
                self.vector_store = Chroma.from_documents(
                    documents=documents,
                    embedding=self.embeddings,
                    persist_directory=self.persist_directory
                )
            else:
                self.vector_store.add_documents(documents)
        except Exception as e:
            # If database is locked/readonly, recreate it
            if "readonly" in str(e).lower() or "locked" in str(e).lower():
                print(f"Database error, recreating: {e}")
                import shutil
                if os.path.exists(self.persist_directory):
                    shutil.rmtree(self.persist_directory)
                os.makedirs(self.persist_directory, exist_ok=True)
                
                self.vector_store = Chroma.from_documents(
                    documents=documents,
                    embedding=self.embeddings,
                    persist_directory=self.persist_directory
                )
            else:
                raise
    
    def add_file(self, file_path: str, file_type: str) -> int:
        """Add a file to knowledge base, returns number of chunks added"""
        documents = self.process_file(file_path, file_type)
        self.add_documents(documents)
        return len(documents)
    
    def query(self, question: str, k: int = 4) -> List[Document]:
        """Query the knowledge base"""
        if self.vector_store is None:
            return []
        
        retriever = self.vector_store.as_retriever(search_kwargs={"k": k})
        return retriever.invoke(question)
    
    def get_all_sources(self) -> List[str]:
        """Get list of all source files in knowledge base"""
        if self.vector_store is None:
            return []
        
        try:
            # Get all documents and extract unique sources
            all_docs = self.vector_store.get()
            if all_docs and 'metadatas' in all_docs:
                sources = set()
                for metadata in all_docs['metadatas']:
                    if 'source' in metadata:
                        sources.add(metadata['source'])
                return list(sources)
        except Exception as e:
            print(f"Error getting sources: {e}")
        
        return []
    
    def clear(self):
        """Clear all documents from knowledge base"""
        if self.vector_store is not None:
            try:
                # Delete the persist directory
                import shutil
                if os.path.exists(self.persist_directory):
                    shutil.rmtree(self.persist_directory)
                os.makedirs(self.persist_directory, exist_ok=True)
                self.vector_store = None
            except Exception as e:
                print(f"Error clearing knowledge base: {e}")

